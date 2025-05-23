from flask import Flask, render_template, request, send_file, jsonify, session
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
from docx2pdf import convert
import requests
import math
import re
import urllib.parse
from PIL import Image, ImageDraw, ImageFont
import base64
import io

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'  # Add a secret key for session
UPLOAD_FOLDER = "output"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 카카오맵 API 키
# KAKAO_API_KEY = "a64f3421ce631504cb1a78f4f4b10036"
KAKAO_API_KEY = os.environ.get('KAKAO_API_KEY')

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/form')
def form():
    form_type = request.args.get('type', 'agriculture')
    my_name = request.args.get('my_name')
    my_address = request.args.get('my_address')
    
    if my_name and my_address:
        session['my_name'] = my_name
        session['my_address'] = my_address
    
    return render_template('form.html', 
                         form_type=form_type,
                         my_name=session.get('my_name'),
                         my_address=session.get('my_address'))

@app.route("/calculate_distance", methods=["POST"])
def calculate_distance():
    data = request.get_json()
    addr1 = data.get("addr1")
    addr2 = data.get("addr2")
    print(f"[거리계산] addr1: {addr1}, addr2: {addr2}")
    
    def get_coordinates(address):
        url = f"https://dapi.kakao.com/v2/local/search/address.json?query={urllib.parse.quote(address)}"
        headers = {
            "Authorization": f"KakaoAK {KAKAO_API_KEY}",
            "User-Agent": "Mozilla/5.0",
            "KA": "python/1.0"
        }
        response = requests.get(url, headers=headers)
        result = response.json()
        print("[카카오API 응답]", result)
        if "documents" in result and result["documents"]:
            return float(result["documents"][0]["y"]), float(result["documents"][0]["x"])
        if "errorMessage" in result and result["errorMessage"] is not None:
            print(f"[카카오API 에러] {str(result["errorMessage"])}")
        return None
    
    def calculate_distance_km(coord1, coord2):
        # Haversine formula
        lat1, lon1 = coord1
        lat2, lon2 = coord2
        
        R = 6371  # Earth's radius in kilometers
        
        lat1, lon1, lat2, lon2 = map(math.radians, [lat1, lon1, lat2, lon2])
        dlat = lat2 - lat1
        dlon = lon2 - lon1
        
        a = math.sin(dlat/2)**2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon/2)**2
        c = 2 * math.asin(math.sqrt(a))
        
        return R * c
    
    coord1 = get_coordinates(addr1)
    coord2 = get_coordinates(addr2)
    
    if not coord1:
        return jsonify({"error": f"출발지 주소를 찾을 수 없습니다: {addr1}"}), 400
    if not coord2:
        return jsonify({"error": f"도착지 주소를 찾을 수 없습니다: {addr2}"}), 400
    
    distance = calculate_distance_km(coord1, coord2)
    return jsonify({"distance": f"{distance:.1f}km"})

def remove_light_bg(input_path, output_path, threshold=220):
    try:
        img = Image.open(input_path).convert("RGBA")
        datas = img.getdata()
        newData = []
        for item in datas:
            brightness = (item[0] + item[1] + item[2]) / 3
            if brightness > threshold:
                newData.append((255, 255, 255, 0))
            else:
                newData.append(item)
        img.putdata(newData)
        img.save(output_path, "PNG")
        return True
    except Exception as e:
        print(f"[오류] 배경 제거 중 오류 발생: {str(e)}")
        return False

def get_korean_font_path():
    candidates = [
        "C:/Windows/Fonts/gulim.ttf",
        "C:/Windows/Fonts/Dotum.ttf",
        "C:/Windows/Fonts/DotumChe.ttf",
        "C:/Windows/Fonts/malgun.ttf",
        "C:/Windows/Fonts/Arial.ttf"
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    raise FileNotFoundError("한글 폰트 파일을 찾을 수 없습니다.")

def get_text_size(draw, text, font):
    # Pillow 최신: draw.textbbox
    if hasattr(draw, "textbbox"):
        bbox = draw.textbbox((0, 0), text, font=font)
        return bbox[2] - bbox[0], bbox[3] - bbox[1]
    # 일부 버전: font.getbbox
    elif hasattr(font, "getbbox"):
        bbox = font.getbbox(text)
        return bbox[2] - bbox[0], bbox[3] - bbox[1]
    # 일부 버전: font.getsize
    elif hasattr(font, "getsize"):
        return font.getsize(text)
    # fallback: (예상치)
    else:
        return (len(text) * 20, 30)

def combine_signature_with_label(signature_path, label_path, output_path):
    from PIL import Image
    # 1. 작업용 캔버스(556x155) 생성
    canvas_w, canvas_h = 556, 155
    canvas = Image.new("RGBA", (canvas_w, canvas_h), (255,255,255,255))
    # 2. 라벨 이미지를 556x155로 리사이즈
    label_img = Image.open(label_path).convert("RGBA")
    label_img = label_img.resize((canvas_w, canvas_h), Image.LANCZOS)
    # 3. 라벨 이미지를 하단에 배치 (여백 0)
    canvas.paste(label_img, (0, 0), label_img)
    # 4. 서명 이미지를 캔버스 폭 100%로 리사이즈, 높이는 라벨 위 공간에 맞춤
    sign_img = Image.open(signature_path).convert("RGBA")
    max_sign_width = canvas_w - 20  # 좌우 10px 여백
    max_sign_height = canvas_h - 20  # 위쪽/아래쪽 여백
    ratio = min(max_sign_width / sign_img.width, max_sign_height / sign_img.height)
    new_sign_size = (int(sign_img.width * ratio), int(sign_img.height * ratio))
    sign_img = sign_img.resize(new_sign_size, Image.LANCZOS)
    sign_x = (canvas_w - sign_img.width) // 2
    sign_y = 10 + (max_sign_height - sign_img.height) // 2
    canvas.paste(sign_img, (sign_x, sign_y), sign_img)
    # 5. 최종 이미지를 232x97로 리사이즈하여 저장
    final_img = canvas.resize((232, 97), Image.LANCZOS)
    final_img.save(output_path, "PNG")

def process_shared_location_image(image_file, target_width_mm, target_height_mm):
    if not image_file:
        return None
        
    try:
        # 이미지 파일 읽기
        img = Image.open(image_file).convert("RGBA") # RGBA로 변환하여 투명도 처리 가능하도록
        
        # 원본 이미지 크기 출력
        print(f"Original image size: {img.size}")
        
        # 템플릿의 목표 크기에 맞춰 조정 (픽셀 계산)
        dpi = 96
        pixels_per_mm = dpi / 25.4 # 1인치 = 25.4mm
        target_width_px = int(target_width_mm * pixels_per_mm)
        target_height_px = int(target_height_mm * pixels_per_mm)
        
        # 가로세로 비율 유지하면서 지정된 칸의 최대 크기에 맞게 리사이즈
        img.thumbnail((target_width_px, target_height_px), Image.Resampling.LANCZOS)
        
        # 리사이즈된 이미지 크기 출력
        print(f"Resized image size: {img.size}")
        
        # 이미지를 바이트로 변환
        img_byte_arr = io.BytesIO()
        # PNG 형식으로 저장 (투명도 지원)
        img.save(img_byte_arr, format='PNG')
        img_byte_arr = img_byte_arr.getvalue()
        
        return img_byte_arr
    except Exception as e:
        print(f"[오류] 공유취득 위치 이미지 처리 중 오류 발생: {str(e)}")
        return None

def insert_shared_location_image(doc, image_data):
    if not image_data or not doc:
        return
        
    try:
        # 이미지 데이터를 임시 파일로 저장
        temp_img = io.BytesIO(image_data)
        
        # 문서에 이미지 추가
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(temp_img, width=Inches(5.9))  # 150mm = 5.9 inches
    except Exception as e:
        print(f"[오류] 공유취득 위치 이미지 삽입 중 오류 발생: {str(e)}")

@app.route('/sample_viewer')
def sample_viewer():
    plan_type_param = request.args.get('type')
    if plan_type_param == 'agriculture':
        plan_type = '농업경영계획서'
    elif plan_type_param == 'weekend':
        plan_type = '주말체험영농계획서'
    else:
        plan_type = '알 수 없는 계획서'

    return render_template('sample_viewer.html', plan_type=plan_type)

@app.route('/generate', methods=['POST'])
def generate():
    form_type = request.form.get('form_type', 'agriculture')
    try:
        print("[시작] 문서 생성 프로세스 시작")
        print(f"[정보] 문서 유형: {form_type}")
        
        # 템플릿 처리
        template_path = "주말체험영농계획서.docx" if form_type == "weekend" else "농업경영계획서.docx"
        print(f"[정보] 템플릿 경로: {template_path}")
        doc = DocxTemplate(template_path)
        print("[성공] 템플릿 로드 완료")
        
        # 기존 데이터 처리
        data = request.form.to_dict()
        data["today"] = datetime.now().strftime("%Y년 %m월 %d일")
        print("[정보] 기본 데이터 처리 완료")
        
        # 공유지분 처리
        for i in range(1, 4):
            area = data.get(f"target_share_area{i}", "").strip()
            denom = data.get(f"target_share_denominator{i}", "").strip()
            numer = data.get(f"target_share_numerator{i}", "").strip()
            if area and denom and numer:
                data[f"target_share{i}"] = f"{area}㎡ 중 {denom}분의 {numer}"
            else:
                data[f"target_share{i}"] = ""
        print("[정보] 공유지분 데이터 처리 완료")
        
        # 날짜 형식 변환
        def ymd_to_kor(val):
            if val and re.match(r"^\d{4}-\d{2}-\d{2}$", val):
                y, m, d = val.split("-")
                return f"{y}년 {int(m)}월 {int(d)}일"
            return val
        data["planting_date"] = ymd_to_kor(data.get("planting_date", ""))
        data["harvest_date"] = ymd_to_kor(data.get("harvest_date", ""))
        print("[정보] 날짜 데이터 처리 완료")
        
        # 서명 처리
        print("[시작] 서명 처리")
        signature_final_path = None
        sign_mode = data.get("sign_mode", "upload")
        
        if sign_mode == "upload" and "signature" in request.files:
            print("[정보] 서명 업로드 모드")
            signature_file = request.files["signature"]
            if signature_file and signature_file.filename:
                try:
                    # 임시 파일 저장
                    signature_path = os.path.join(UPLOAD_FOLDER, "temp_signature.png")
                    signature_file.save(signature_path)
                    print("[성공] 서명 파일 저장 완료")
                    
                    # 배경 제거
                    signature_nobg_path = os.path.join(UPLOAD_FOLDER, "temp_signature_nobg.png")
                    if remove_light_bg(signature_path, signature_nobg_path):
                        print("[성공] 서명 배경 제거 완료")
                        # 라벨과 결합
                        label_path = os.path.join("static", "sign_label.png")
                        combined_path = os.path.join(UPLOAD_FOLDER, "temp_signature_combined.png")
                        combine_signature_with_label(signature_nobg_path, label_path, combined_path)
                        signature_final_path = combined_path
                        print("[성공] 서명 라벨 결합 완료")
                    else:
                        print("[경고] 배경 제거 실패, 원본 이미지 사용")
                        signature_final_path = signature_path
                except Exception as e:
                    print(f"[오류] 서명 처리 중 오류 발생: {str(e)}")
                    return f"서명 처리 중 오류가 발생했습니다: {str(e)}", 500
                    
        elif sign_mode == "draw" and data.get("signature_draw"):
            print("[정보] 서명 그리기 모드")
            try:
                img_data = data["signature_draw"].split(",")[1]
                signature_path = os.path.join(UPLOAD_FOLDER, "temp_signature_draw.png")
                with open(signature_path, "wb") as f:
                    f.write(base64.b64decode(img_data))
                print("[성공] 서명 그리기 파일 저장 완료")
                
                signature_nobg_path = os.path.join(UPLOAD_FOLDER, "temp_signature_draw_nobg.png")
                if remove_light_bg(signature_path, signature_nobg_path):
                    print("[성공] 서명 그리기 배경 제거 완료")
                    label_path = os.path.join("static", "sign_label.png")
                    combined_path = os.path.join(UPLOAD_FOLDER, "temp_signature_draw_combined.png")
                    combine_signature_with_label(signature_nobg_path, label_path, combined_path)
                    signature_final_path = combined_path
                    print("[성공] 서명 그리기 라벨 결합 완료")
                else:
                    print("[경고] 배경 제거 실패, 원본 이미지 사용")
                    signature_final_path = signature_path
            except Exception as e:
                print(f"[오류] 서명 그리기 처리 중 오류 발생: {str(e)}")
                return f"서명 그리기 처리 중 오류가 발생했습니다: {str(e)}", 500
        
        # 서명 추가
        if signature_final_path:
            data["signature"] = InlineImage(doc, signature_final_path, width=Mm(25))
            print("[성공] 서명 이미지 추가 완료")
        
        # 공유취득 위치 이미지 처리
        print("[시작] 공유취득 위치 이미지 처리")
        if 'shared_location' in request.files:
            # 템플릿의 {{shared_location}} 변수 위치에 삽입될 이미지의 목표 크기 (mm)
            target_shared_location_width_mm = 150 # 추정치
            target_shared_location_height_mm = 60 # 추정치
            
            shared_location_img = process_shared_location_image(request.files['shared_location'], target_shared_location_width_mm, target_shared_location_height_mm)
            if shared_location_img:
                # 이미지를 임시 파일로 저장
                temp_img_path = os.path.join(UPLOAD_FOLDER, "temp_shared_location.png")
                with open(temp_img_path, "wb") as f:
                    f.write(shared_location_img)
                # InlineImage로 변환하여 템플릿에 추가
                data["shared_location"] = InlineImage(doc, temp_img_path, width=Mm(target_shared_location_width_mm), height=Mm(target_shared_location_height_mm))
                print("[성공] 공유취득 위치 이미지 추가 완료")
            else:
                # 이미지가 없거나 처리 실패 시 빈 이미지 생성
                blank_img = Image.new('RGBA', (int(target_shared_location_width_mm * 3.78), int(target_shared_location_height_mm * 3.78)), (255, 255, 255, 0))
                temp_img_path = os.path.join(UPLOAD_FOLDER, "temp_blank_shared_location.png")
                blank_img.save(temp_img_path, "PNG")
                data["shared_location"] = InlineImage(doc, temp_img_path, width=Mm(target_shared_location_width_mm), height=Mm(target_shared_location_height_mm))
                print("[정보] 빈 이미지로 대체 완료")
        
        # 문서 저장
        print("[시작] 문서 저장")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_path = os.path.join(UPLOAD_FOLDER, f"plan_{form_type}_{timestamp}.docx")
        print(f"[정보] 저장 경로: {docx_path}")
        doc.render(data)
        doc.save(docx_path)
        print("[성공] DOCX 파일 저장 완료")
        
        # PDF 변환 또는 DOCX 반환
        if request.form.get("output_type") == "pdf":
            print("[시작] PDF 변환")
            pdf_path = os.path.join(UPLOAD_FOLDER, f"plan_{form_type}_{timestamp}.pdf")
            import pythoncom
            pythoncom.CoInitialize()
            convert(docx_path, pdf_path)
            print("[성공] PDF 변환 완료")
            return send_file(pdf_path, as_attachment=True)
        else:
            return send_file(docx_path, as_attachment=True)
            
    except Exception as e:
        print(f"[오류] 문서 생성 중 오류 발생: {str(e)}")
        import traceback
        print(f"[오류 상세] {traceback.format_exc()}")
        return f"문서 생성 중 오류가 발생했습니다: {str(e)}", 500

if __name__ == "__main__":
    app.run(debug=True)
